"""
Document extraction — extract text from uploaded files.

Supports PDF, DOCX, CSV, TXT, and common code files.
Dependencies are optional: gracefully degrades if PyMuPDF or python-docx
are not installed.

Inspired by openclaw's file-processing pipeline.
"""
from __future__ import annotations

import csv
import io
import logging
import mimetypes
import os
from dataclasses import dataclass
from pathlib import Path
from typing import Optional

log = logging.getLogger(__name__)

MAX_FILE_SIZE = int(os.getenv("EXTRACT_MAX_SIZE_MB", "50")) * 1024 * 1024
MAX_TEXT_LENGTH = int(os.getenv("EXTRACT_MAX_TEXT_CHARS", "500000"))

# Code file extensions we can read as plain text
_CODE_EXTENSIONS = {
    ".py", ".js", ".ts", ".jsx", ".tsx", ".java", ".c", ".cpp", ".h",
    ".hpp", ".cs", ".go", ".rs", ".rb", ".php", ".swift", ".kt",
    ".scala", ".sh", ".bash", ".zsh", ".fish", ".ps1",
    ".sql", ".r", ".m", ".mm", ".lua", ".pl", ".pm",
    ".yaml", ".yml", ".toml", ".ini", ".cfg", ".conf",
    ".xml", ".html", ".htm", ".css", ".scss", ".less",
    ".json", ".jsonl", ".ndjson",
    ".md", ".rst", ".tex", ".txt", ".log",
    ".dockerfile", ".makefile", ".cmake",
    ".gitignore", ".env.example", ".editorconfig",
}


@dataclass
class ExtractResult:
    text: str
    pages: int
    format: str
    error: str | None = None
    truncated: bool = False


def _check_deps() -> dict[str, bool]:
    """Check which optional extraction dependencies are available."""
    deps = {}
    for mod in ("fitz", "docx"):
        try:
            __import__(mod)
            deps[mod] = True
        except ImportError:
            deps[mod] = False
    return deps


def available_formats() -> list[str]:
    """Return list of supported file formats."""
    formats = ["txt", "csv", "code files"]
    deps = _check_deps()
    if deps.get("fitz"):
        formats.append("pdf")
    if deps.get("docx"):
        formats.append("docx")
    return formats


def extract_pdf(path: str) -> ExtractResult:
    """Extract text from PDF using PyMuPDF (fitz)."""
    try:
        import fitz
    except ImportError:
        return ExtractResult(
            text="", pages=0, format="pdf",
            error="PyMuPDF not installed. Install with: pip install PyMuPDF",
        )

    try:
        doc = fitz.open(path)
        pages = []
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text()
            if text.strip():
                pages.append(f"--- Page {page_num + 1} ---\n{text}")

        full_text = "\n\n".join(pages)
        truncated = False
        if len(full_text) > MAX_TEXT_LENGTH:
            full_text = full_text[:MAX_TEXT_LENGTH] + "\n\n[...truncated...]"
            truncated = True

        doc.close()
        return ExtractResult(
            text=full_text,
            pages=len(pages),
            format="pdf",
            truncated=truncated,
        )
    except Exception as e:
        return ExtractResult(text="", pages=0, format="pdf", error=str(e)[:200])


def extract_docx(path: str) -> ExtractResult:
    """Extract text from DOCX using python-docx."""
    try:
        import docx
    except ImportError:
        return ExtractResult(
            text="", pages=0, format="docx",
            error="python-docx not installed. Install with: pip install python-docx",
        )

    try:
        doc = docx.Document(path)
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

        full_text = "\n\n".join(paragraphs)
        truncated = False
        if len(full_text) > MAX_TEXT_LENGTH:
            full_text = full_text[:MAX_TEXT_LENGTH] + "\n\n[...truncated...]"
            truncated = True

        return ExtractResult(
            text=full_text,
            pages=len(paragraphs),  # approximate "sections"
            format="docx",
            truncated=truncated,
        )
    except Exception as e:
        return ExtractResult(text="", pages=0, format="docx", error=str(e)[:200])


def extract_csv(path: str) -> ExtractResult:
    """Extract text from CSV files."""
    try:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            # Sniff the dialect
            sample = f.read(8192)
            f.seek(0)
            try:
                dialect = csv.Sniffer().sniff(sample)
            except csv.Error:
                dialect = csv.excel

            reader = csv.reader(f, dialect)
            rows = []
            for i, row in enumerate(reader):
                if i >= 10000:  # safety limit
                    rows.append("[...truncated at 10000 rows...]")
                    break
                rows.append(" | ".join(row))

        full_text = "\n".join(rows)
        truncated = False
        if len(full_text) > MAX_TEXT_LENGTH:
            full_text = full_text[:MAX_TEXT_LENGTH] + "\n\n[...truncated...]"
            truncated = True

        return ExtractResult(
            text=full_text,
            pages=len(rows),
            format="csv",
            truncated=truncated,
        )
    except Exception as e:
        return ExtractResult(text="", pages=0, format="csv", error=str(e)[:200])


def extract_text_file(path: str) -> ExtractResult:
    """Extract text from plain text or code files."""
    try:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read(MAX_TEXT_LENGTH + 1)

        truncated = len(text) > MAX_TEXT_LENGTH
        if truncated:
            text = text[:MAX_TEXT_LENGTH] + "\n\n[...truncated...]"

        line_count = text.count("\n") + 1
        ext = Path(path).suffix.lower()
        fmt = ext.lstrip(".") if ext else "txt"

        return ExtractResult(
            text=text,
            pages=line_count,
            format=fmt,
            truncated=truncated,
        )
    except Exception as e:
        return ExtractResult(text="", pages=0, format="txt", error=str(e)[:200])


def extract(path: str) -> ExtractResult:
    """Extract text from a file, auto-detecting format.

    Args:
        path: Path to the file to extract from.

    Returns:
        ExtractResult with the extracted text, page count, and format.
    """
    if not os.path.exists(path):
        return ExtractResult(text="", pages=0, format="unknown",
                             error=f"File not found: {path}")

    file_size = os.path.getsize(path)
    if file_size > MAX_FILE_SIZE:
        return ExtractResult(
            text="", pages=0, format="unknown",
            error=f"File too large ({file_size // 1024 // 1024}MB, max {MAX_FILE_SIZE // 1024 // 1024}MB)",
        )

    if file_size == 0:
        return ExtractResult(text="", pages=0, format="empty",
                             error="File is empty")

    ext = Path(path).suffix.lower()

    if ext == ".pdf":
        return extract_pdf(path)
    elif ext == ".docx":
        return extract_docx(path)
    elif ext == ".csv":
        return extract_csv(path)
    elif ext in _CODE_EXTENSIONS or ext in (".txt", ".log", ".md", ".rst"):
        return extract_text_file(path)
    else:
        # Try reading as text anyway
        try:
            return extract_text_file(path)
        except Exception:
            return ExtractResult(
                text="", pages=0, format=ext.lstrip(".") or "unknown",
                error=f"Unsupported file format: {ext or 'no extension'}",
            )


def summarize_extraction(result: ExtractResult) -> str:
    """Format extraction result for display."""
    if result.error:
        return f"Error extracting {result.format}: {result.error}"

    lines = [f"Extracted from {result.format.upper()} ({result.pages} pages/sections)"]
    if result.truncated:
        lines.append("(truncated due to size)")

    preview = result.text[:500]
    if len(result.text) > 500:
        preview += "..."
    lines.append(f"\n{preview}")
    return "\n".join(lines)
